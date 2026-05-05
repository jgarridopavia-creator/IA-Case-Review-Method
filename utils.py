"""
Utilidades del Ágora: extracción de texto de PDFs y exportación a Word.
"""
from io import BytesIO
from datetime import datetime
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def extraer_texto_pdf(archivo) -> str:
    """Extrae el texto de un PDF subido vía Streamlit (UploadedFile)."""
    try:
        archivo.seek(0)
        lector = PdfReader(archivo)
        partes = []
        for i, pagina in enumerate(lector.pages, start=1):
            txt = pagina.extract_text() or ""
            if txt.strip():
                partes.append(f"[Página {i}]\n{txt.strip()}")
        return "\n\n".join(partes)
    except Exception as e:
        return f"[ERROR al leer PDF: {e}]"


def extraer_texto_pdfs(archivos) -> str:
    """Extrae y concatena el texto de varios PDFs."""
    if not archivos:
        return ""
    bloques = []
    for archivo in archivos:
        nombre = getattr(archivo, "name", "documento.pdf")
        contenido = extraer_texto_pdf(archivo)
        bloques.append(f"=== DOCUMENTO: {nombre} ===\n{contenido}")
    return "\n\n".join(bloques)


def contar_palabras(texto: str) -> int:
    return len((texto or "").split())


# ---------------------------------------------------------------------------
# Exportación a Word
# ---------------------------------------------------------------------------

def _add_titulo(doc, texto, nivel=1):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    if nivel == 0:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif nivel == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x2E, 0x59, 0x84)


def _add_parrafo(doc, texto):
    if not texto:
        texto = "(sin contenido)"
    for linea in str(texto).split("\n"):
        p = doc.add_paragraph(linea)
        for run in p.runs:
            run.font.size = Pt(11)


def construir_docx(
    nombre_caso: str,
    pregunta: str,
    respuesta_inicial: str,
    respuesta_critico: str,
    respuesta_optimista: str,
    respuesta_final: str,
    respuesta_asertivo: str,
) -> BytesIO:
    """Genera el .docx con toda la sesión y devuelve un BytesIO listo para descargar."""
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Portada
    _add_titulo(doc, "ÁGORA — Análisis de Caso Práctico MBA", nivel=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Caso analizado: {nombre_caso or '—'}")
    run.italic = True
    run.font.size = Pt(11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    _add_titulo(doc, "1. Pregunta planteada por el alumno", nivel=1)
    _add_parrafo(doc, pregunta)

    _add_titulo(doc, "2. Respuesta inicial del alumno", nivel=1)
    _add_parrafo(doc, respuesta_inicial)

    _add_titulo(doc, "3. Análisis del Agente Crítico", nivel=1)
    _add_parrafo(doc, respuesta_critico)

    _add_titulo(doc, "4. Análisis del Agente Optimista", nivel=1)
    _add_parrafo(doc, respuesta_optimista)

    _add_titulo(doc, "5. Respuesta final del alumno", nivel=1)
    _add_parrafo(doc, respuesta_final)

    _add_titulo(doc, "6. Corrección del Agente Asertivo", nivel=1)
    _add_parrafo(doc, respuesta_asertivo)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
